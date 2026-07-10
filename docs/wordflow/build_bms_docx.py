from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
WORD_DIR = ROOT / "docs" / "wordflow"
ASSET_DIR = WORD_DIR / "assets"
MD_PATH = WORD_DIR / "bms24v_project_document_v0.md"
OUT_PATH = WORD_DIR / "bms24v_project_document_v0.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 46)
MUTED = RGBColor(91, 103, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
TODO_FILL = "FFF4CC"
RISK_FILL = "FDECEC"
BORDER = "A7B1BD"
BODY_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"


def set_run_font(run, size=None, color=None, bold=None, italic=None, name=BODY_FONT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_inline_markdown(paragraph, text, size=11, color=INK):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size - 0.5, color=color, name="Consolas")
            run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_para(doc, text, style=None, size=11, color=INK, after=6, before=0, bold=False):
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p, before=before, after=after, line=1.25)
    if bold:
        run = p.add_run(text)
        set_run_font(run, size=size, color=color, bold=True)
    else:
        add_inline_markdown(p, text, size=size, color=color)
    return p


def add_callout(doc, text, fill=LIGHT_GRAY, label=None):
    p = doc.add_paragraph()
    set_paragraph_shading(p, fill)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    set_paragraph_spacing(p, before=4, after=8, line=1.2)
    if label:
        run = p.add_run(label + " ")
        set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
    add_inline_markdown(p, text, size=10.5, color=INK)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    set_paragraph_shading(p, "F7F7F7")
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    set_paragraph_spacing(p, before=4, after=8, line=1.15)
    run = p.add_run(text)
    set_run_font(run, size=8.5, color=INK, name="Consolas")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def clean_cell(text):
    return text.strip().replace("<br>", "\n").replace("\\|", "|")


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        line = lines[idx].strip()
        cells = [clean_cell(c) for c in line.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in cells):
            rows.append(cells)
        idx += 1
    if not rows:
        return None, start + 1
    max_cols = max(len(r) for r in rows)
    for row in rows:
        row.extend([""] * (max_cols - len(row)))
    return rows, idx


def widths_for_table(rows):
    headers = rows[0]
    n = len(headers)
    if n == 1:
        return [9360]
    if n == 2:
        if any("证据" in h or "Evidence" in h for h in headers):
            return [2300, 7060]
        return [2700, 6660]
    if n == 3:
        return [2100, 3400, 3860]
    if n == 4:
        return [1400, 3100, 2100, 2760]
    if n == 5:
        return [1200, 2600, 3500, 860, 1200]
    if n == 6:
        return [760, 1850, 2500, 2500, 1250, 500]
    base = 9360 // n
    widths = [base] * n
    widths[-1] += 9360 - sum(widths)
    return widths


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    widths = widths_for_table(rows)
    set_table_width(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line=1.15)
            add_inline_markdown(p, text, size=8.2 if len(rows[0]) >= 5 else 8.8, color=INK)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
    mark_header_row(table.rows[0])
    doc.add_paragraph()


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header_p, before=0, after=0, line=1.0)
    run = header_p.add_run("24V/6S BMS 项目初版文档 | wordflow")
    set_run_font(run, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_p.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(footer_p)
    run = footer_p.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def generate_diagram_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    try:
        font = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 26)
        small = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 20)
    except Exception:
        font = ImageFont.load_default()
        small = ImageFont.load_default()

    def box(draw, xy, text, fill="#E8EEF5", outline="#5B6B7A"):
        draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
        lines = text.split("\n")
        y = xy[1] + (xy[3] - xy[1] - len(lines) * 30) / 2
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            x = xy[0] + (xy[2] - xy[0] - (bbox[2] - bbox[0])) / 2
            draw.text((x, y), line, fill="#14202E", font=font)
            y += 34

    def arrow(draw, start, end, color="#1F4D78"):
        draw.line([start, end], fill=color, width=4)
        ex, ey = end
        sx, sy = start
        dx, dy = ex - sx, ey - sy
        if abs(dx) > abs(dy):
            pts = [(ex, ey), (ex - 18 if dx > 0 else ex + 18, ey - 10), (ex - 18 if dx > 0 else ex + 18, ey + 10)]
        else:
            pts = [(ex, ey), (ex - 10, ey - 18 if dy > 0 else ey + 18), (ex + 10, ey - 18 if dy > 0 else ey + 18)]
        draw.polygon(pts, fill=color)

    img = Image.new("RGB", (1800, 980), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 36), "硬件-软件总体关系图（说明图，非原理图）", fill="#1F4D78", font=font)
    box(d, (90, 220, 350, 340), "6S 三元锂\n电芯")
    box(d, (430, 210, 730, 350), "BQ76952\n采样/保护/均衡")
    box(d, (830, 210, 1160, 350), "STM32G0B1CBT6\nApp / Int / Com")
    box(d, (1260, 180, 1610, 320), "SC8815\n升降压充电")
    box(d, (1260, 410, 1610, 530), "24V_IN / VBUS")
    box(d, (430, 500, 730, 620), "主 FET\nCHG/DSG/PDSG")
    box(d, (830, 500, 1160, 620), "FDCAN / OLED\nEEPROM / CLI")
    box(d, (1260, 650, 1610, 770), "BMS+ / 电池包")
    arrow(d, (350, 280), (430, 280))
    arrow(d, (730, 280), (830, 280))
    arrow(d, (1160, 280), (1260, 250))
    arrow(d, (1435, 410), (1435, 320))
    arrow(d, (1435, 530), (1435, 650))
    arrow(d, (580, 350), (580, 500))
    arrow(d, (730, 560), (830, 560))
    d.text((96, 860), "依据：.ioc、Int_BQ76952/Int_SC8815、App_Power、网表/PDF。", fill="#5B6770", font=small)
    img.save(ASSET_DIR / "bms24v_architecture_logic.png")

    img = Image.new("RGB", (1800, 980), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 36), "App_Power 状态机说明图（代码阅读辅助）", fill="#1F4D78", font=font)
    positions = {
        "OFF": (130, 220, 380, 330),
        "BQ_WAKE": (520, 120, 830, 240),
        "RUN": (520, 380, 830, 500),
        "LOW": (950, 260, 1220, 380),
        "MONITOR": (950, 520, 1260, 640),
        "FAULT": (1340, 380, 1630, 500),
    }
    labels = {
        "OFF": "OFF\n输出关闭",
        "BQ_WAKE": "BQ_WAKE\nSC 唤醒 BQ",
        "RUN": "RUN\n允许充/放电",
        "LOW": "LOW\n低压/过流",
        "MONITOR": "MONITOR\n恢复观察",
        "FAULT": "FAULT\n故障锁存",
    }
    for key, xy in positions.items():
        fill = RISK_FILL if key == "FAULT" else LIGHT_BLUE
        box(d, xy, labels[key], fill="#" + fill if not fill.startswith("#") else fill)
    arrow(d, (380, 275), (520, 180))
    arrow(d, (380, 275), (520, 440))
    arrow(d, (675, 240), (675, 380))
    arrow(d, (830, 440), (950, 320))
    arrow(d, (1085, 380), (1105, 520))
    arrow(d, (1260, 580), (1340, 455))
    arrow(d, (830, 440), (1340, 440))
    for text, xy in [
        ("BQ离线且有24V输入", (410, 150)),
        ("电芯有效且无故障", (410, 455)),
        ("BQ恢复采样", (690, 300)),
        ("低压或放电过流", (790, 295)),
        ("RC补偿后恢复", (980, 455)),
        ("SCD/故障", (1080, 410)),
    ]:
        d.text(xy, text, fill="#5B6770", font=small)
    d.text((96, 860), "依据：App/App_Power.c:321-522。状态图不代表实测结果。", fill="#5B6770", font=small)
    img.save(ASSET_DIR / "bms24v_power_state_machine.png")


def add_title_block(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=4, line=1.1)
    run = p.add_run("24V/6S BMS 项目文档")
    set_run_font(run, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=14, line=1.2)
    run = p.add_run("基于当前 new_bms 代码、网表、PDF 与旧版 UPS 文档写法整理")
    set_run_font(run, size=12, color=MUTED)

    rows = [
        ("版本", "v0.3 / Word 初版"),
        ("分支", "wordflow"),
        ("范围", "STM32G0B1CBT6、BQ76952、SC8815、FDCAN、OLED、EEPROM、Debug CLI"),
        ("规则", "按旧 Word 教学节奏写；事实按当前证据落地；缺图/缺实测写 TODO"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [1550, 7810])
    for i, (k, v) in enumerate(rows):
        for j, text in enumerate((k, v)):
            cell = table.cell(i, j)
            set_cell_margins(cell, top=90, bottom=90, start=140, end=140)
            if j == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line=1.15)
            add_inline_markdown(p, text, size=9.5, color=INK)
            if j == 0:
                p.runs[0].bold = True
    mark_header_row(table.rows[0])
    doc.add_paragraph()


def collect_toc(lines):
    headings = []
    for line in lines:
        m = re.match(r"^(#{2,3})\s+(.+)", line.strip())
        if m:
            headings.append((len(m.group(1)) - 1, m.group(2)))
    return headings


def add_static_toc(doc, lines):
    p = doc.add_heading("目录", level=1)
    p.paragraph_format.page_break_before = False
    for level, text in collect_toc(lines):
        if level > 1 and not re.match(r"\d+\.\d+", text):
            continue
        p = doc.add_paragraph(style="List Bullet" if level == 1 else "List Bullet 2")
        set_paragraph_spacing(p, before=0, after=2, line=1.15)
        add_inline_markdown(p, text, size=10, color=INK)


def image_path_from_markdown(line):
    m = re.match(r"!\[(.*?)\]\((.*?)\)", line.strip())
    if not m:
        return None, None
    alt, rel = m.group(1), m.group(2)
    return alt, (WORD_DIR / rel).resolve()


def add_image(doc, image_path, alt):
    if not image_path.exists():
        add_callout(doc, f"//TODO 此处应该放 {alt}，但图片文件缺失：`{image_path}`。", fill=TODO_FILL, label="TODO")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.2))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=8, line=1.1)
    run = p.add_run(alt)
    set_run_font(run, size=9.2, color=MUTED, italic=True)


def build_docx():
    generate_diagram_assets()
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_title_block(doc)
    add_callout(
        doc,
        "本文按旧 UPS 项目文档的讲义节奏组织：先讲项目和硬件，再讲工程配置、接口层代码和 APP 业务。未实测、未截图、未确认的内容保留 Unknown / Conflict / //TODO。",
        fill=LIGHT_GRAY,
        label="阅读说明",
    )
    add_static_toc(doc, lines)
    doc.add_page_break()

    i = 0
    in_code = False
    code_lang = ""
    code_lines = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("# "):
            i += 1
            continue
        if stripped.startswith("版本：") or stripped.startswith("分支：") or stripped.startswith("范围：") or stripped.startswith("写作规则："):
            i += 1
            continue

        if in_code:
            if stripped == "```":
                text = "\n".join(code_lines)
                if code_lang == "mermaid":
                    if "flowchart" in text:
                        add_image(doc, ASSET_DIR / "bms24v_architecture_logic.png", "硬件-软件总体关系图（说明图）")
                    elif "stateDiagram" in text:
                        add_image(doc, ASSET_DIR / "bms24v_power_state_machine.png", "App_Power 状态机说明图")
                    else:
                        add_code_block(doc, text)
                else:
                    add_code_block(doc, text)
                in_code = False
                code_lang = ""
                code_lines = []
            else:
                code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            code_lang = stripped.strip("`").strip()
            code_lines = []
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        img_alt, img_path = image_path_from_markdown(stripped)
        if img_path is not None:
            add_image(doc, img_path, img_alt)
            i += 1
            continue

        if stripped.startswith("|"):
            rows, new_i = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
                i = new_i
                continue

        if stripped.startswith("//TODO"):
            add_callout(doc, stripped, fill=TODO_FILL, label="TODO")
            i += 1
            continue

        m = re.match(r"^(#{2,4})\s+(.+)", stripped)
        if m:
            md_level = len(m.group(1))
            text = m.group(2)
            level = 1 if md_level == 2 else 2 if md_level == 3 else 3
            doc.add_heading(text, level=level)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            set_paragraph_spacing(p, before=0, after=4, line=1.25)
            add_inline_markdown(p, re.sub(r"^\d+\.\s+", "", stripped), size=10.5, color=INK)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, before=0, after=4, line=1.25)
            add_inline_markdown(p, stripped[2:], size=10.5, color=INK)
            i += 1
            continue

        if stripped.startswith("结论：Conflict") or stripped.startswith("Unknown："):
            add_callout(doc, stripped, fill=RISK_FILL if "Conflict" in stripped else TODO_FILL, label="重点")
            i += 1
            continue

        if stripped.startswith("说明："):
            add_callout(doc, stripped, fill=LIGHT_GRAY, label="说明")
            i += 1
            continue

        add_para(doc, stripped, size=10.8, color=INK)
        i += 1

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_docx()
    print(OUT_PATH)
