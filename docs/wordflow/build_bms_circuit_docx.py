from __future__ import annotations

import csv
import math
import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_bms_docx as base


ROOT = Path(__file__).resolve().parents[2]
WORD_DIR = ROOT / "docs" / "wordflow"
ASSET_DIR = WORD_DIR / "assets"
MD_PATH = WORD_DIR / "bms_hardware_circuit_design_audit_v1.md"
OUT_PATH = WORD_DIR / "new_bms_24V6S_硬件电路设计说明书_v1.0.docx"

# documents skill design contract:
# - one preset only: compact_reference_guide
# - first-page pattern: editorial_cover
# - the existing project palette below is one named override, not a second preset.
DESIGN_PRESET = "compact_reference_guide"
FIRST_PAGE_PATTERN = "editorial_cover"
COLOR_OVERRIDE = "new_bms_evidence_palette"

PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
PAGE_MARGIN_IN = 1.0
HEADER_FOOTER_DISTANCE_IN = 0.492
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_DXA = 80
CELL_MARGIN_BOTTOM_DXA = 80
CELL_MARGIN_START_DXA = 120
CELL_MARGIN_END_DXA = 120

BODY_SIZE_PT = 11
BODY_AFTER_PT = 6
BODY_LINE_SPACING = 1.25
COVER_KICKER_SIZE_PT = 10
COVER_TITLE_SIZE_PT = 27
COVER_SUBTITLE_SIZE_PT = 12.5
# Named component overrides within the same preset (dense evidence tables/captions).
TABLE_TEXT_SIZE_PT = 8.8
WIDE_TABLE_TEXT_SIZE_PT = 8.2
CAPTION_TEXT_SIZE_PT = 9.2
HEADING_TOKENS = {
    "Heading 1": (16, 18, 10),
    "Heading 2": (13, 14, 7),
    "Heading 3": (12, 10, 5),
}
LIST_MARKER_DXA = 270
LIST_TEXT_DXA = 540
LIST_HANGING_DXA = 270
LIST_AFTER_TWIPS = 80
LIST_LINE_TWIPS = 300

# This builder produces the delivery candidate. In final mode, absent figures are
# evidence failures and therefore abort the build instead of becoming placeholders.
FINAL_MODE = True

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(91, 103, 112)
INK = RGBColor(20, 31, 46)

APPENDIX_ACTIVE_WIDTHS = [650, 800, 1750, 4360, 1800]
APPENDIX_NET_WIDTHS = [1800, 900, 4260, 2400]
MIN_TABLE_COLUMN_DXA = 620
NET_ENDPOINTS_PER_RECORD = 24

# All body close-ups are regenerated directly from the original vector PDFs at
# 320 DPI.  Embed the Markdown-referenced PNG itself: redirecting it to an
# appendix image would apply stale whole-page pixel coordinates to an already
# cropped image and could silently remove electrically meaningful content.
SCHEMATIC_CROPS = {}

# Some high-resolution close-ups still contain a partial neighbouring panel or
# a large empty margin. These normalized bounds remove only that residue. Very
# wide regions retain a complete overview and then use named, overlapping
# semantic detail panels; medium-width figures stay intact.
SCHEMATIC_VIEW_BOUNDS = {
    "bms_p01_connectors.png": (0.0, 0.0, 1.0, 0.72),
    "bms_p01_bq_control.png": (0.0, 0.0, 1.0, 0.80),
    "bms_p01_cell_balance.png": (0.0, 0.02, 1.0, 0.84),
    "control_p05_clock_reset_swd.png": (0.0, 0.0, 1.0, 0.605),
}
SCHEMATIC_SEMANTIC_PANELS = {
    "bms_p01_connectors.png": (
        ((0.010, 0.07, 0.184, 0.72), "平衡线接口 CN2"),
        ((0.183, 0.07, 0.408, 0.72), "XT90 功率线接口 CN1"),
        ((0.407, 0.07, 0.657, 0.72), "24V 电源管理跨板接口 U26"),
        ((0.656, 0.07, 0.944, 0.72), "M3 定位件 U22-U25"),
    ),
    "bms_p01_main_power_path.png": (
        ((0.00, 0.00, 0.34, 1.00), "BAT+、CHG 驱动与 Q5/Q6 充电 MOS 组"),
        ((0.26, 0.00, 0.56, 1.00), "共漏节点、Q1/Q2 放电 MOS 组与 DSG 驱动"),
        ((0.48, 0.00, 0.79, 1.00), "Q8 加速关断与 Q3/D19 反接辅助网络"),
        ((0.70, 0.00, 1.00, 1.00), "BMS_OUT+ 汇出、钳位与端子"),
    ),
    "control_p01_sc_power_stage.png": (
        ((0.00, 0.00, 0.44, 1.00), "VBUS 输入、SNS1 分流与 Q14/Q2 左桥臂"),
        ((0.35, 0.00, 0.76, 1.00), "L1 与 Q14/Q2、Q15/Q16 四开关核心"),
        ((0.65, 0.00, 1.00, 1.00), "Q15/Q16 右桥臂、SNS2 分流、去耦、D2 与 BMS+"),
    ),
    "control_p05_clock_reset_swd.png": (
        ((0.005, 0.02, 0.220, 0.604), "8MHz HSE 晶振"),
        ((0.219, 0.02, 0.452, 0.604), "32.768kHz LSE 晶振"),
        ((0.451, 0.02, 0.684, 0.604), "NRST 上拉、RC 与复位按键"),
        ((0.683, 0.02, 0.989, 0.604), "SWD、NRST 与调试 UART 接口"),
    ),
}
INDEPENDENT_SEMANTIC_PANELS = {
    "bms_p01_connectors.png",
    "control_p05_clock_reset_swd.png",
}
PARTIAL_CONTEXT_FIGURES = {
    "bms_p01_bq_control.png",
    "bms_p01_cell_balance.png",
}
MAX_SCHEMATIC_PANEL_ASPECT = 4.0
MAX_SCHEMATIC_PANEL_COUNT = 4
SCHEMATIC_PANEL_OVERLAP = 0.18
SCHEMATIC_PANEL_WIDTH_IN = 6.2
FRONT_MATTER_PREFIXES = (
    "版本：",
    "日期：",
    "对象：",
    "证据截止：",
    "范围：",
    "写作规则：",
)


def _set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        successors = {
            qn(tag)
            for tag in (
                "w:hdrShapeDefaults",
                "w:footnotePr",
                "w:endnotePr",
                "w:compat",
                "w:docVars",
                "w:rsids",
                "w:themeFontLang",
                "w:clrSchemeMapping",
                "w:doNotIncludeSubdocsInStats",
                "w:doNotAutoCompressPictures",
                "w:forceUpgrade",
                "w:captions",
                "w:readModeInkLockDown",
                "w:smartTagType",
                "w:shapeDefaults",
                "w:doNotEmbedSmartTags",
                "w:decimalSymbol",
                "w:listSeparator",
                "w14:docId",
                "w14:defaultImageDpi",
            )
        }
        first_successor = next((child for child in settings if child.tag in successors), None)
        if first_successor is None:
            settings.append(update)
        else:
            settings.insert(settings.index(first_successor), update)
    update.set(qn("w:val"), "true")


def _next_id(nodes, attr: str) -> int:
    values = []
    for node in nodes:
        raw = node.get(qn(attr))
        if raw is not None and raw.isdigit():
            values.append(int(raw))
    return max(values, default=0) + 1


def _insert_abstract_num(numbering, abstract) -> None:
    """Keep the numbering part schema-valid: every abstractNum precedes w:num."""
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        num_id_mac = numbering.find(qn("w:numIdMac"))
        if num_id_mac is None:
            numbering.append(abstract)
        else:
            numbering.insert(numbering.index(num_id_mac), abstract)
        return
    numbering.insert(numbering.index(first_num), abstract)


def _insert_num(numbering, num) -> None:
    """Insert concrete numbering before optional numIdMac, per CT_Numbering order."""
    num_id_mac = numbering.find(qn("w:numIdMac"))
    if num_id_mac is None:
        numbering.append(num)
    else:
        numbering.insert(numbering.index(num_id_mac), num)


def _make_numbering(doc: Document, kind: str) -> int:
    """Create real compact-reference numbering definitions."""
    if LIST_TEXT_DXA - LIST_HANGING_DXA != LIST_MARKER_DXA:
        raise ValueError("列表 marker/text/hanging token 彼此不一致")
    numbering = doc.part.numbering_part.element
    abstract_id = _next_id(numbering.findall(qn("w:abstractNum")), "w:abstractNumId")
    num_id = _next_id(numbering.findall(qn("w:num")), "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    bullet_text = ["•", "○", "▪"]
    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), bullet_text[level] if kind == "bullet" else f"%{level + 1}.")
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        text_indent = LIST_TEXT_DXA * (level + 1)
        tab.set(qn("w:pos"), str(text_indent))
        tabs.append(tab)
        p_pr.append(tabs)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), str(LIST_AFTER_TWIPS))
        spacing.set(qn("w:line"), str(LIST_LINE_TWIPS))
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(text_indent))
        ind.set(qn("w:hanging"), str(LIST_HANGING_DXA))
        p_pr.append(ind)
        lvl.append(p_pr)
        if kind == "bullet":
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), base.BODY_FONT)
            r_fonts.set(qn("w:hAnsi"), base.BODY_FONT)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)

    _insert_abstract_num(numbering, abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    _insert_num(numbering, num)
    return num_id


def _restart_numbering_instance(
    doc: Document,
    template_num_id: int,
    *,
    start: int = 1,
    level: int = 0,
) -> int:
    """Create a concrete numbering instance for one Markdown list block."""
    if start < 1 or not 0 <= level <= 2:
        raise ValueError("列表起始值或层级无效")

    numbering = doc.part.numbering_part.element
    template = next(
        (
            node
            for node in numbering.findall(qn("w:num"))
            if node.get(qn("w:numId")) == str(template_num_id)
        ),
        None,
    )
    if template is None:
        raise ValueError(f"未找到编号模板 numId={template_num_id}")
    abstract_ref = template.find(qn("w:abstractNumId"))
    if abstract_ref is None or abstract_ref.get(qn("w:val")) is None:
        raise ValueError(f"编号模板 numId={template_num_id} 缺少 abstractNumId")

    num_id = _next_id(numbering.findall(qn("w:num")), "w:numId")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), abstract_ref.get(qn("w:val")))
    num.append(abstract_num_id)

    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), str(level))
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start))
    level_override.append(start_override)
    num.append(level_override)
    _insert_num(numbering, num)
    return num_id


def _apply_numbering(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def _set_style_font(style, *, size: float, color: RGBColor, bold: bool) -> None:
    style.font.name = base.BODY_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), base.BODY_FONT)
    r_fonts.set(qn("w:hAnsi"), base.BODY_FONT)
    r_fonts.set(qn("w:eastAsia"), base.CJK_FONT)


def _apply_compact_reference_tokens(doc: Document) -> None:
    """Resolve compact_reference_guide into explicit, auditable Word values."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(PAGE_MARGIN_IN)
    section.right_margin = Inches(PAGE_MARGIN_IN)
    section.bottom_margin = Inches(PAGE_MARGIN_IN)
    section.left_margin = Inches(PAGE_MARGIN_IN)
    section.header_distance = Inches(HEADER_FOOTER_DISTANCE_IN)
    section.footer_distance = Inches(HEADER_FOOTER_DISTANCE_IN)

    normal = doc.styles["Normal"]
    _set_style_font(normal, size=BODY_SIZE_PT, color=INK, bold=False)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(BODY_AFTER_PT)
    normal.paragraph_format.line_spacing = BODY_LINE_SPACING

    for name, (size, before, after) in HEADING_TOKENS.items():
        color = BLUE if name != "Heading 3" else NAVY
        style = doc.styles[name]
        _set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = BODY_LINE_SPACING
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    # Word's built-in TOC Heading uses outline level 9 (body text), so the TOC
    # cannot recursively include its own title. Its visual tokens match H1.
    toc_heading = doc.styles["TOC Heading"]
    h1_size, _, h1_after = HEADING_TOKENS["Heading 1"]
    _set_style_font(toc_heading, size=h1_size, color=BLUE, bold=True)
    toc_heading.paragraph_format.space_before = Pt(0)
    toc_heading.paragraph_format.space_after = Pt(h1_after)
    toc_heading.paragraph_format.line_spacing = BODY_LINE_SPACING
    toc_heading.paragraph_format.keep_with_next = True
    toc_heading.paragraph_format.keep_together = True


def configure_document(doc: Document) -> tuple[int, int]:
    base.configure_document(doc)
    _apply_compact_reference_tokens(doc)
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    for run in list(header_p.runs):
        run._element.getparent().remove(run._element)
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    base.set_paragraph_spacing(header_p, before=0, after=0, line=1.0)
    run = header_p.add_run("new_bms 24V/6S 硬件电路设计说明书 | 证据版")
    base.set_run_font(run, size=9, color=MUTED)

    doc.core_properties.title = "new_bms 24V/6S 硬件电路设计说明书"
    doc.core_properties.subject = "基于原理图、网表、数据手册与固件的电路逻辑说明"
    doc.core_properties.author = "Codex（证据化生成）"
    doc.core_properties.keywords = (
        "BMS, BQ76952, SC8815, STM32G0B1, LM7480-Q1, netlist, "
        f"preset:{DESIGN_PRESET}, cover:{FIRST_PAGE_PATTERN}, palette:{COLOR_OVERRIDE}"
    )
    _set_update_fields(doc)
    return _make_numbering(doc, "bullet"), _make_numbering(doc, "decimal")


def add_title_block(doc: Document) -> None:
    """Apply the editorial_cover first-page pattern for the technical guide."""
    spacer = doc.add_paragraph()
    base.set_paragraph_spacing(spacer, before=72, after=0, line=1.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_paragraph_spacing(p, before=0, after=16, line=1.0)
    run = p.add_run("硬件设计说明书 · EVIDENCE EDITION")
    base.set_run_font(run, size=COVER_KICKER_SIZE_PT, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_paragraph_spacing(p, before=0, after=10, line=1.05)
    run = p.add_run("new_bms 24V/6S\n硬件电路设计与软硬件映射说明书")
    base.set_run_font(run, size=COVER_TITLE_SIZE_PT, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_paragraph_spacing(p, before=0, after=52, line=1.25)
    run = p.add_run("逐页原理图解析 · 全网表互证 · 官方数据手册校核 · 固件引脚与安全时序对齐")
    base.set_run_font(run, size=COVER_SUBTITLE_SIZE_PT, color=MUTED)

    rows = [
        ("文档版本", "v1.0 / 2026-07-14"),
        ("硬件范围", "6S BQ76952 BMS 板 + STM32G0B1/SC8815 24V 电源管理板"),
        ("主要证据", "11 页原理图、2 份完整网表、官方芯片资料、当前固件与 .ioc"),
        ("结论规则", "事实/推断/Unknown/Conflict 分离；安全相关项须经实板验证"),
    ]
    # These are definitions, not columnar records. Keeping them as labeled
    # paragraphs avoids falsely marking "文档版本" as a repeating table header.
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.55)
        p.paragraph_format.right_indent = Inches(0.55)
        p.paragraph_format.keep_together = True
        base.set_paragraph_spacing(p, before=0, after=4, line=1.15)
        label_run = p.add_run(f"{label}：")
        base.set_run_font(label_run, size=9.5, color=BLUE, bold=True)
        base.add_inline_markdown(p, value, size=9.5, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_paragraph_spacing(p, before=34, after=0, line=1.15)
    run = p.add_run("本说明书不把尚未实测的电气行为写成已验证事实。")
    base.set_run_font(run, size=10.5, color=RGBColor(155, 28, 28), bold=True)
    doc.add_page_break()


def collect_toc(lines: list[str]) -> list[tuple[int, str]]:
    """Collect Markdown H2/H3/H4 as Word Heading 1/2/3 TOC entries."""
    headings = []
    in_code = False
    for line in lines:
        stripped = line.strip()
        if not in_code and stripped.startswith("```"):
            in_code = True
            continue
        if in_code:
            if stripped == "```":
                in_code = False
            continue
        if stripped == "[[APPENDIX:ACTIVE_COMPONENTS]]":
            headings.append((1, "附录 A：有源器件、功率器件与连接器逐引脚索引"))
            continue
        if stripped == "[[APPENDIX:NAMED_NETS]]":
            headings.append((1, "附录 B：命名网络连接索引"))
            continue
        match = re.match(r"^(#{2,4})\s+(.+)", stripped)
        if match:
            headings.append((len(match.group(1)) - 1, match.group(2).strip()))
    return headings


def _append_field_char(run, field_type: str, *, dirty: bool = False) -> None:
    node = OxmlElement("w:fldChar")
    node.set(qn("w:fldCharType"), field_type)
    if dirty:
        node.set(qn("w:dirty"), "true")
    run._r.append(node)


def add_toc(doc: Document, lines: list[str]) -> None:
    """Insert a real, updateable Heading 1-3 TOC with cached heading text."""
    title = doc.add_paragraph(style="TOC Heading")
    run = title.add_run("目录")
    base.set_run_font(run, size=16, color=BLUE, bold=True)

    headings = collect_toc(lines)
    if not headings:
        headings = [(1, "打开 Word 后更新目录字段")]

    for index, (level, text) in enumerate(headings):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18 + (level - 1) * 0.30)
        p.paragraph_format.keep_together = True
        base.set_paragraph_spacing(p, before=0, after=2, line=1.12)
        if index == 0:
            field_run = p.add_run()
            _append_field_char(field_run, "begin", dirty=True)
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = ' TOC \\o "1-3" \\h \\z \\u '
            field_run._r.append(instr)
            _append_field_char(field_run, "separate")

        base.add_inline_markdown(p, text, size=10 if level == 1 else 9.3, color=INK)
        if index == len(headings) - 1:
            end_run = p.add_run()
            _append_field_char(end_run, "end")

    doc.add_page_break()


def _parse_image(line: str) -> tuple[str | None, Path | None, float]:
    match = re.match(r"!\[(.*?)\]\((.*?)\)(?:\{width=([0-9.]+)\})?", line.strip())
    if not match:
        return None, None, 6.2
    alt, rel, raw_width = match.groups()
    image_path = (WORD_DIR / rel).resolve()
    if not image_path.is_relative_to(ROOT.resolve()):
        raise ValueError(f"图片路径越出项目目录：{image_path}")
    width = float(raw_width or 6.2)
    if width <= 0.0 or width > 6.2:
        raise ValueError(f"图片宽度必须在 (0, 6.2] 英寸：{image_path} -> {width}")
    return alt, image_path, width


def _png_dimensions(image_path: Path) -> tuple[int, int] | None:
    """Read PNG dimensions without adding an image-processing dependency."""
    with image_path.open("rb") as stream:
        header = stream.read(24)
    if len(header) < 24 or header[:8] != b"\x89PNG\r\n\x1a\n" or header[12:16] != b"IHDR":
        return None
    width = int.from_bytes(header[16:20], "big")
    height = int.from_bytes(header[20:24], "big")
    if width <= 0 or height <= 0:
        return None
    return width, height


def _schematic_base_bounds(
    image_path: Path,
    source_size: tuple[int, int],
    pixel_crop: tuple[int, int, int, int] | None,
) -> tuple[float, float, float, float]:
    """Return the normalized visible region before any detail-panel split."""
    source_width, source_height = source_size
    if pixel_crop is not None:
        left_px, top_px, right_px, bottom_px = pixel_crop
        if not (
            0 <= left_px < right_px <= source_width
            and 0 <= top_px < bottom_px <= source_height
        ):
            raise ValueError(f"图片裁剪越界：{image_path} -> {pixel_crop}")
        base_bounds = (
            left_px / source_width,
            top_px / source_height,
            right_px / source_width,
            bottom_px / source_height,
        )
    else:
        base_bounds = SCHEMATIC_VIEW_BOUNDS.get(image_path.name, (0.0, 0.0, 1.0, 1.0))

    left, top, right, bottom = base_bounds
    if not (0.0 <= left < right <= 1.0 and 0.0 <= top < bottom <= 1.0):
        raise ValueError(f"图片裁切范围无效：{image_path.name} -> {base_bounds}")
    return base_bounds


def _schematic_panel_bounds(
    image_path: Path,
    source_size: tuple[int, int],
    pixel_crop: tuple[int, int, int, int] | None,
) -> list[tuple[float, float, float, float]]:
    """Return named semantic panels or a conservative overlapping fallback."""
    source_width, source_height = source_size
    base_bounds = _schematic_base_bounds(image_path, source_size, pixel_crop)
    left, top, right, bottom = base_bounds

    semantic_panels = SCHEMATIC_SEMANTIC_PANELS.get(image_path.name)
    if semantic_panels is not None:
        bounds_list = [bounds for bounds, _ in semantic_panels]
        for bounds in bounds_list:
            panel_left, panel_top, panel_right, panel_bottom = bounds
            if not (
                left <= panel_left < panel_right <= right
                and top <= panel_top < panel_bottom <= bottom
            ):
                raise ValueError(
                    f"语义分区越出可见区域：{image_path.name} -> {bounds} / {base_bounds}"
                )
        return bounds_list

    visible_aspect = (source_width * (right - left)) / (source_height * (bottom - top))
    panel_count = min(
        MAX_SCHEMATIC_PANEL_COUNT,
        max(1, math.ceil(visible_aspect / MAX_SCHEMATIC_PANEL_ASPECT)),
    )
    if panel_count == 1:
        return [base_bounds]

    span = (right - left) / panel_count
    overlap = span * SCHEMATIC_PANEL_OVERLAP
    panels = []
    for panel_index in range(panel_count):
        panel_left = left + panel_index * span
        panel_right = left + (panel_index + 1) * span
        if panel_index > 0:
            panel_left -= overlap
        if panel_index < panel_count - 1:
            panel_right += overlap
        panels.append((max(left, panel_left), top, min(right, panel_right), bottom))
    return panels


def _apply_picture_crop(shape, bounds: tuple[float, float, float, float]) -> None:
    """Apply DrawingML crop values while preserving the visible aspect ratio."""
    left, top, right, bottom = bounds
    visible_width = right - left
    visible_height = bottom - top
    if visible_width <= 0.0 or visible_height <= 0.0:
        raise ValueError(f"图片可见区域无效：{bounds}")

    original_width = int(shape.width)
    original_height = int(shape.height)
    blip_fill = shape._inline.graphic.graphicData.pic.blipFill
    src_rect = blip_fill.find(qn("a:srcRect"))
    if src_rect is None:
        src_rect = OxmlElement("a:srcRect")
        stretch = blip_fill.find(qn("a:stretch"))
        if stretch is None:
            blip_fill.append(src_rect)
        else:
            blip_fill.insert(blip_fill.index(stretch), src_rect)
    for attribute, fraction in (
        ("l", left),
        ("t", top),
        ("r", 1.0 - right),
        ("b", 1.0 - bottom),
    ):
        src_rect.set(attribute, str(round(fraction * 100000)))
    shape.height = round(
        original_width
        * (original_height / original_width)
        * (visible_height / visible_width)
    )


def _add_picture_paragraph(
    doc: Document,
    display_path: Path,
    *,
    width: float,
    bounds: tuple[float, float, float, float],
    description: str,
) -> None:
    """Insert one inline picture and bind it to the immediately following caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    base.set_paragraph_spacing(p, before=0, after=2, line=1.0)
    run = p.add_run()
    shape = run.add_picture(str(display_path), width=Inches(width))
    _apply_picture_crop(shape, bounds)
    shape._inline.docPr.set("descr", description)
    shape._inline.docPr.set("title", description)


def _add_figure_caption(doc: Document, text: str, *, after: float = 8) -> None:
    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_together = True
    caption.paragraph_format.keep_with_next = False
    caption.paragraph_format.widow_control = True
    base.set_paragraph_spacing(caption, before=1, after=after, line=1.1)
    run = caption.add_run(text)
    base.set_run_font(run, size=CAPTION_TEXT_SIZE_PT, color=MUTED, italic=True)


def add_image(
    doc: Document,
    image_path: Path,
    alt: str,
    width: float,
    *,
    final_mode: bool = FINAL_MODE,
) -> None:
    alt = alt.strip()
    if final_mode and not alt:
        raise ValueError(f"最终模式下图片必须有替代文本：{image_path}")

    display_path = image_path
    crop_box: tuple[int, int, int, int] | None = None
    crop_spec = SCHEMATIC_CROPS.get(image_path.name)
    if crop_spec is not None:
        source_name, crop_box, width_override = crop_spec
        display_path = image_path.with_name(source_name)
        if width_override is not None:
            width = width_override

    if not display_path.exists():
        if final_mode:
            raise FileNotFoundError(f"最终模式下图片缺失，已终止构建：{display_path}")
        base.add_callout(
            doc,
            f"图片缺失：`{display_path}`。",
            fill=base.TODO_FILL,
            label="Unknown",
        )
        return

    source_size = _png_dimensions(display_path)
    if source_size is None:
        base_bounds = (0.0, 0.0, 1.0, 1.0)
        panels = [base_bounds]
    else:
        base_bounds = _schematic_base_bounds(image_path, source_size, crop_box)
        panels = _schematic_panel_bounds(image_path, source_size, crop_box)

    if len(panels) == 1:
        description = alt
        caption_text = alt
        if image_path.name in PARTIAL_CONTEXT_FIGURES:
            caption_text += "（原始矢量局部放大；图边外连接见附录整页原理图）"
            description = caption_text
        _add_picture_paragraph(
            doc,
            display_path,
            width=width,
            bounds=panels[0],
            description=description,
        )
        _add_figure_caption(doc, caption_text)
        return

    independent_panels = image_path.name in INDEPENDENT_SEMANTIC_PANELS
    overview_alt = f"{alt}（完整总览）"
    _add_picture_paragraph(
        doc,
        display_path,
        width=width,
        bounds=base_bounds,
        description=overview_alt,
    )
    if independent_panels:
        overview_caption = f"{alt}（完整总览；下列原始功能框分别放大）"
    else:
        overview_caption = f"{alt}（完整总览；下列分区从左到右放大，相邻分区保留重叠）"
    _add_figure_caption(doc, overview_caption, after=5)

    semantic_panels = SCHEMATIC_SEMANTIC_PANELS.get(image_path.name)
    for panel_index, bounds in enumerate(panels, start=1):
        if semantic_panels is None:
            panel_label = "连续网络局部"
        else:
            panel_label = semantic_panels[panel_index - 1][1]
        panel_alt = f"{alt}（放大分区 {panel_index}/{len(panels)}：{panel_label}）"
        _add_picture_paragraph(
            doc,
            display_path,
            width=SCHEMATIC_PANEL_WIDTH_IN,
            bounds=bounds,
            description=panel_alt,
        )
        if independent_panels:
            caption_text = f"功能框 {panel_index}/{len(panels)} — {panel_label}"
        else:
            caption_text = (
                f"放大分区 {panel_index}/{len(panels)} — {panel_label}（与相邻分区有重叠）"
            )
        _add_figure_caption(
            doc,
            caption_text,
            after=5 if panel_index < len(panels) else 8,
        )


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    """Apply fixed DXA geometry: tblW/tblInd/tblGrid/tcW must agree."""
    if not widths_dxa or sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(
            f"表格列宽必须合计 {CONTENT_WIDTH_DXA} DXA，实际为 {sum(widths_dxa)}"
        )
    if len(widths_dxa) != len(table.columns):
        raise ValueError("表格列宽数量与列数不一致")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert_element_before(tbl_w, "w:jc", "w:tblCellSpacing", "w:tblInd")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.insert_element_before(
            tbl_ind,
            "w:tblBorders",
            "w:shd",
            "w:tblLayout",
            "w:tblCellMar",
            "w:tblLook",
        )
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.insert_element_before(layout, "w:tblCellMar", "w:tblLook")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(table._tbl.index(tbl_pr) + 1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for column, cell in enumerate(row.cells):
            base.set_cell_width(cell, widths_dxa[column])


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def add_table(
    doc: Document,
    rows: list[list[str]],
    widths_dxa: list[int] | None = None,
    *,
    keep_rows_together: bool = False,
):
    """Create a semantic header table using the compact-reference table tokens."""
    if not rows or not rows[0]:
        raise ValueError("不能创建空表格")
    column_count = len(rows[0])
    if any(len(row) != column_count for row in rows):
        raise ValueError("表格各行列数不一致")
    widths = list(widths_dxa or _widths_for_table(rows))

    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index, text in enumerate(row):
            cell = table.cell(row_index, column_index)
            if row_index == 0:
                base.set_cell_shading(cell, base.LIGHT_BLUE)
            base.set_cell_margins(
                cell,
                top=CELL_MARGIN_TOP_DXA,
                bottom=CELL_MARGIN_BOTTOM_DXA,
                start=CELL_MARGIN_START_DXA,
                end=CELL_MARGIN_END_DXA,
            )
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            base.set_paragraph_spacing(p, before=0, after=0, line=1.15)
            base.add_inline_markdown(
                p,
                text,
                size=WIDE_TABLE_TEXT_SIZE_PT if column_count >= 5 else TABLE_TEXT_SIZE_PT,
                color=INK,
            )
            if row_index == 0:
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.keep_together = True
                for run in p.runs:
                    run.bold = True

    _set_table_geometry(table, widths)
    base.mark_header_row(table.rows[0])
    _set_row_cant_split(table.rows[0])
    # A split evidence/FMEA row is no longer a verifiable record: Word repeats
    # only some cells on the next page and visually creates a false blank row.
    # Keep every semantic record intact; rows are never given a fixed height.
    for row in table.rows[1:]:
        _set_row_cant_split(row)
    spacer = doc.add_paragraph()
    base.set_paragraph_spacing(spacer, before=0, after=2, line=1.0)
    return table


def _table_header_weight(header: str) -> float:
    """Map a header's semantic role to its relative reading width."""
    value = re.sub(r"[`*_]", "", header).strip().lower()
    if any(
        token in value
        for token in ("必须决策", "必须验证", "通过条件", "通过标准", "所需证据", "验证/改进")
    ):
        return 2.75
    if "evidence" in value or "证据" in value or value.startswith("source"):
        return 2.35
    if "human confirmation" in value or "人工确认" in value:
        return 1.95
    if any(token in value for token in ("结论", "conclusion", "预期系统结果")):
        return 2.45
    if any(
        token in value
        for token in (
            "设计职责",
            "设计含义",
            "工程解释",
            "可能后果",
            "软件行为",
            "软件动作",
            "软件控制",
            "主要职责",
            "当前所有者",
            "硬件作用",
            "不能声称",
            "使用范围",
            "证据边界",
            "支持的结论",
            "操作",
            "作用/限制",
            "安全影响",
        )
    ):
        return 2.05
    if any(
        token in value
        for token in ("路径", "网络", "连接", "器件", "外围", "配置", "分压", "失效模式", "unknown")
    ):
        return 1.55
    if any(token in value for token in ("confidence", "风险", "判定", "分类", "状态")):
        return 1.10
    if any(
        token in value
        for token in ("id", "序号", "步骤", "cell", "pin", "引脚", "地址", "优先级", "栈", "页")
    ):
        return 0.82
    if any(
        token in value
        for token in ("电流", "电压", "功耗", "门限", "速率", "标称", "写值", "结果", "范围")
    ):
        return 1.12
    return 1.35


def _widths_for_table(rows: list[list[str]]) -> list[int]:
    """Allocate the exact 9360 DXA width from semantic header roles."""
    headers = rows[0]
    column_count = len(headers)
    if column_count == 1:
        return [CONTENT_WIDTH_DXA]
    minimum_total = MIN_TABLE_COLUMN_DXA * column_count
    if minimum_total >= CONTENT_WIDTH_DXA:
        raise ValueError(f"表格列数过多，无法满足最小列宽：{column_count}")

    weights = [_table_header_weight(header) for header in headers]
    distributable = CONTENT_WIDTH_DXA - minimum_total
    weight_total = sum(weights)
    raw_extras = [distributable * weight / weight_total for weight in weights]
    widths = [MIN_TABLE_COLUMN_DXA + int(extra) for extra in raw_extras]
    remainder = CONTENT_WIDTH_DXA - sum(widths)
    order = sorted(
        range(column_count),
        key=lambda column: raw_extras[column] - int(raw_extras[column]),
        reverse=True,
    )
    for column in order[:remainder]:
        widths[column] += 1
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"语义列宽合计错误：{widths}")
    return widths


def _split_markdown_row(line: str) -> list[str]:
    """Split a pipe row while preserving Markdown's escaped ``\\|`` content."""
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        backslashes = 0
        for char in reversed(raw[:-1]):
            if char != "\\":
                break
            backslashes += 1
        if backslashes % 2 == 0:
            raw = raw[:-1]

    cells: list[str] = []
    current: list[str] = []
    index = 0
    while index < len(raw):
        char = raw[index]
        if char == "\\" and index + 1 < len(raw) and raw[index + 1] == "|":
            current.append("|")
            index += 2
            continue
        if char == "|":
            cells.append("".join(current).strip().replace("<br>", "\n"))
            current = []
        else:
            current.append(char)
        index += 1
    cells.append("".join(current).strip().replace("<br>", "\n"))
    return cells


def parse_table(lines: list[str], start: int):
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = _split_markdown_row(lines[index])
        if not all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells):
            rows.append(cells)
        index += 1
    if not rows:
        return None, start + 1
    max_columns = max(len(row) for row in rows)
    for row in rows:
        row.extend([""] * (max_columns - len(row)))
    return rows, index


def _audit_table_geometry(doc: Document) -> None:
    for table_index, table in enumerate(doc.tables, start=1):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        layout = tbl_pr.find(qn("w:tblLayout"))
        grid_widths = [
            int(node.get(qn("w:w"))) for node in table._tbl.tblGrid.findall(qn("w:gridCol"))
        ]
        if table.autofit:
            raise ValueError(f"表 {table_index} 未关闭 autofit")
        if layout is None or layout.get(qn("w:type")) != "fixed":
            raise ValueError(f"表 {table_index} 缺少 fixed tblLayout")
        if (
            tbl_w is None
            or tbl_w.get(qn("w:type")) != "dxa"
            or int(tbl_w.get(qn("w:w"), "-1")) != CONTENT_WIDTH_DXA
            or sum(grid_widths) != CONTENT_WIDTH_DXA
        ):
            raise ValueError(f"表 {table_index} 的 tblW/tblGrid 不一致")
        if (
            tbl_ind is None
            or tbl_ind.get(qn("w:type")) != "dxa"
            or int(tbl_ind.get(qn("w:w"), "-1")) != TABLE_INDENT_DXA
        ):
            raise ValueError(f"表 {table_index} 的 tblInd 不符合预设")
        for row_index, row in enumerate(table.rows, start=1):
            tr_pr = row._tr.find(qn("w:trPr"))
            if tr_pr is None or tr_pr.find(qn("w:cantSplit")) is None:
                raise ValueError(f"表 {table_index} 第 {row_index} 行允许跨页拆分")
            if row_index == 1:
                for column_index, cell in enumerate(row.cells, start=1):
                    for paragraph in cell.paragraphs:
                        p_pr = paragraph._p.find(qn("w:pPr"))
                        if p_pr is None or p_pr.find(qn("w:keepNext")) is None:
                            raise ValueError(
                                f"表 {table_index} 表头第 {column_index} 列缺少 keepNext"
                            )
            for column_index, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if (
                    tc_w is None
                    or tc_w.get(qn("w:type")) != "dxa"
                    or int(tc_w.get(qn("w:w"), "-1")) != grid_widths[column_index]
                ):
                    raise ValueError(
                        f"表 {table_index} 第 {row_index} 行 tcW 与 tblGrid 不一致"
                    )


def _natural_key(value: str):
    return [int(part) if part.isdigit() else part.lower() for part in re.split(r"(\d+)", value)]


def _description_summary(description: str) -> str:
    fields = []
    for key in (
        "阻值",
        "容值",
        "额定电压",
        "功率",
        "型号",
        "漏源电压(Vdss)",
        "导通电阻(RDS(on))",
        "保持电流(Ihold)",
        "跳闸电流(Itrip)",
    ):
        match = re.search(rf"{re.escape(key)}:([^;]+)", description)
        if match:
            fields.append(f"{key}={match.group(1).strip()}")
    if fields:
        return "; ".join(fields[:4])
    return description.strip()[:100] or "Unknown"


def _read_netlist(path: Path):
    with path.open("r", encoding="utf-8-sig", newline="") as stream:
        rows = list(csv.DictReader(stream))
    for line_no, row in enumerate(rows, start=2):
        row["_line"] = line_no
    return rows


def _compress_line_refs(lines: list[int], max_span: int = 4) -> str:
    """Return exact CSV line references, splitting long runs into small ranges."""
    ordered = sorted(set(lines))
    if not ordered:
        return "Unknown"
    ranges: list[str] = []
    run_start = ordered[0]
    run_end = ordered[0]

    def flush(start: int, end: int) -> None:
        cursor = start
        while cursor <= end:
            chunk_end = min(cursor + max_span - 1, end)
            ranges.append(f"L{cursor}" if cursor == chunk_end else f"L{cursor}-L{chunk_end}")
            cursor = chunk_end + 1

    for line in ordered[1:]:
        if line == run_end + 1:
            run_end = line
            continue
        flush(run_start, run_end)
        run_start = run_end = line
    flush(run_start, run_end)
    return ",".join(ranges)


def _netlist_evidence(filename: str, lines: list[int]) -> str:
    return f"official_chip_docs_files/{filename}:{_compress_line_refs(lines)}"


def _include_active_appendix_record(designator: str, category: str) -> bool:
    keep_categories = {
        "IC/模块",
        "三极管",
        "接插件",
        "连接器",
        "保险丝",
        "电感",
        "晶振",
        "开关",
        "二极管",
        "LED",
        "蜂鸣器",
    }
    if category in keep_categories:
        return True
    if category != "其他":
        return False
    # The two netlists place several real electromechanical/active records in
    # “其他”; use reference prefixes so these are not silently discarded.
    return bool(re.match(r"^(?:BUZZER|CN|FH|F|H|J|LDO|LED|SW|X)\d", designator.upper()))


def add_active_component_appendix(doc: Document) -> None:
    doc.add_heading("附录 A：有源器件、功率器件与连接器逐引脚索引", level=1)
    base.add_callout(
        doc,
        "本附录由两份当前网表机械生成，保留自动网络名和未命名引脚；描述字段只能证明网表元数据，不能替代器件官方手册。",
        fill=base.LIGHT_GRAY,
        label="范围",
    )
    for filename, board in (
        ("full_netlist (4).csv", "BQ76952 BMS 板"),
        ("full_netlist (5).csv", "STM32/SC8815 控制板"),
    ):
        path = ROOT / "official_chip_docs_files" / filename
        rows = _read_netlist(path)
        grouped = defaultdict(list)
        for row in rows:
            grouped[row["Designator"]].append(row)
        records = []
        for designator, items in grouped.items():
            category = items[0]["Category"]
            if not _include_active_appendix_record(designator, category):
                continue
            pins = []
            for item in items:
                net = item["Net"] or "NC/未命名"
                pin_name = item["PinName"].strip()
                label = f"{item['Pin']}({pin_name})" if pin_name else item["Pin"]
                pins.append(f"{label}={net}")
            lines = [item["_line"] for item in items]
            evidence = _netlist_evidence(filename, lines)
            records.append(
                [
                    designator,
                    category,
                    _description_summary(items[0]["Description"]),
                    "; ".join(pins),
                    evidence,
                ]
            )
        records.sort(key=lambda row: _natural_key(row[0]))
        doc.add_heading(board, level=2)
        chunk_size = 18
        for start in range(0, len(records), chunk_size):
            add_table(
                doc,
                [["位号", "类别", "关键参数", "逐引脚网络", "网表证据"], *records[start : start + chunk_size]],
                APPENDIX_ACTIVE_WIDTHS,
            )


def add_named_net_appendix(doc: Document) -> None:
    doc.add_heading("附录 B：命名网络连接索引", level=1)
    base.add_callout(
        doc,
        "仅列出具有明确名称且至少连接两个引脚的网络；大网按每块最多 24 个端点拆为“网络（1/n）”，每块独立保留精确 CSV 行证据。`$1N.../$2N.../$3N...` 自动网络仍保留在附录 A 的逐引脚表中。",
        fill=base.LIGHT_GRAY,
        label="范围",
    )
    for filename, board in (
        ("full_netlist (4).csv", "BQ76952 BMS 板"),
        ("full_netlist (5).csv", "STM32/SC8815 控制板"),
    ):
        path = ROOT / "official_chip_docs_files" / filename
        rows = _read_netlist(path)
        nets = defaultdict(list)
        for row in rows:
            net = row["Net"].strip()
            if not net or net.startswith("$"):
                continue
            nets[net].append(row)
        sortable_records = []
        for net, items in nets.items():
            if len(items) < 2:
                continue
            ordered_items = sorted(items, key=lambda item: item["_line"])
            total = len(ordered_items)
            part_count = (total + NET_ENDPOINTS_PER_RECORD - 1) // NET_ENDPOINTS_PER_RECORD
            for part_index, start in enumerate(
                range(0, total, NET_ENDPOINTS_PER_RECORD),
                start=1,
            ):
                part_items = ordered_items[start : start + NET_ENDPOINTS_PER_RECORD]
                endpoints = []
                for item in part_items:
                    pin_name = item["PinName"].strip()
                    endpoint = f"{item['Designator']}.{item['Pin']}"
                    if pin_name:
                        endpoint += f"({pin_name})"
                    endpoints.append(endpoint)
                display_net = net if part_count == 1 else f"{net}（{part_index}/{part_count}）"
                sortable_records.append(
                    (
                        net,
                        part_index,
                        [
                            display_net,
                            f"{len(part_items)} / {total}",
                            "; ".join(endpoints),
                            _netlist_evidence(filename, [item["_line"] for item in part_items]),
                        ],
                    )
                )
        sortable_records.sort(key=lambda record: (_natural_key(record[0]), record[1]))
        records = [record[2] for record in sortable_records]
        doc.add_heading(board, level=2)
        for start in range(0, len(records), 22):
            add_table(
                doc,
                [["网络", "本块/总数", "连接端点", "网表证据"], *records[start : start + 22]],
                APPENDIX_NET_WIDTHS,
                keep_rows_together=True,
            )


def _validate_code_fences(lines: list[str]) -> None:
    opening_line: int | None = None
    for line_number, line in enumerate(lines, start=1):
        stripped = line.strip()
        if opening_line is None and stripped.startswith("```"):
            opening_line = line_number
        elif opening_line is not None and stripped == "```":
            opening_line = None
    if opening_line is not None:
        raise ValueError(f"Markdown 代码块在 EOF 前未闭合（起始行 L{opening_line}）")


def build_docx(*, final_mode: bool = FINAL_MODE, output_path: Path = OUT_PATH) -> Path:
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    _validate_code_fences(lines)
    doc = Document()
    bullet_num_id, decimal_num_id = configure_document(doc)
    add_title_block(doc)
    base.add_callout(
        doc,
        "正文采用“功能目的 -> 电流/信号路径 -> 元件作用 -> 控制条件 -> 软件映射 -> 计算与风险”的统一读法。每一项安全结论都带证据等级；没有资料或没有实测的项目保持 Unknown。",
        fill=base.LIGHT_GRAY,
        label="阅读方法",
    )
    add_toc(doc, lines)

    index = 0
    in_code = False
    code_lines: list[str] = []
    code_start_line: int | None = None
    in_front_matter = True
    active_decimal_num_id: int | None = None
    last_decimal_marker: int | None = None
    last_decimal_level: int | None = None
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        numbered_line = None if in_code else re.match(r"^(\s*)(\d+)\.\s+(.+)", line)
        if stripped and numbered_line is None:
            active_decimal_num_id = None
            last_decimal_marker = None
            last_decimal_level = None
        if in_front_matter:
            if not stripped or stripped.startswith("# ") or stripped.startswith(FRONT_MATTER_PREFIXES):
                index += 1
                continue
            in_front_matter = False
        if stripped.startswith("# "):
            index += 1
            continue

        if in_code:
            if stripped == "```":
                base.add_code_block(doc, "\n".join(code_lines))
                in_code = False
                code_lines = []
                code_start_line = None
            else:
                code_lines.append(line)
            index += 1
            continue
        if stripped.startswith("```"):
            in_code = True
            code_lines = []
            code_start_line = index + 1
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped == "<!-- PAGEBREAK -->":
            doc.add_page_break()
            index += 1
            continue
        if stripped == "[[APPENDIX:ACTIVE_COMPONENTS]]":
            add_active_component_appendix(doc)
            index += 1
            continue
        if stripped == "[[APPENDIX:NAMED_NETS]]":
            add_named_net_appendix(doc)
            index += 1
            continue

        alt, image_path, width = _parse_image(stripped)
        if image_path is not None and alt is not None:
            add_image(doc, image_path, alt, width, final_mode=final_mode)
            index += 1
            continue

        if stripped.startswith("|"):
            rows, new_index = parse_table(lines, index)
            if rows:
                add_table(doc, rows)
                index = new_index
                continue

        callout_match = re.match(r"> \[(FACT|INFERENCE|CALC|RISK|UNKNOWN|CONFLICT|EVIDENCE)\]\s*(.+)", stripped)
        if callout_match:
            kind, text = callout_match.groups()
            fill = {
                "FACT": base.LIGHT_BLUE,
                "INFERENCE": "F4F6F9",
                "CALC": "EEF6EE",
                "RISK": "FDECEC",
                "UNKNOWN": base.TODO_FILL,
                "CONFLICT": "FDECEC",
                "EVIDENCE": "F4F6F9",
            }[kind]
            base.add_callout(doc, text, fill=fill, label=kind.title())
            index += 1
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)", stripped)
        if heading:
            md_level = len(heading.group(1))
            level = 1 if md_level == 2 else 2 if md_level == 3 else 3
            p = doc.add_heading(heading.group(2), level=level)
            p.paragraph_format.keep_with_next = True
            index += 1
            continue

        if numbered_line:
            level = min(len(numbered_line.group(1)) // 2, 2)
            source_marker = int(numbered_line.group(2))
            starts_new_list = active_decimal_num_id is None or (
                source_marker == 1
                and last_decimal_marker is not None
                and last_decimal_level is not None
                and level <= last_decimal_level
            )
            if starts_new_list:
                active_decimal_num_id = _restart_numbering_instance(
                    doc,
                    decimal_num_id,
                    start=source_marker,
                    level=level,
                )
            p = doc.add_paragraph()
            _apply_numbering(p, active_decimal_num_id, level)
            base.set_paragraph_spacing(p, before=0, after=4, line=1.25)
            base.add_inline_markdown(p, numbered_line.group(3), size=BODY_SIZE_PT, color=INK)
            last_decimal_marker = source_marker
            last_decimal_level = level
            index += 1
            continue

        bullet = re.match(r"^(\s*)-\s+(.+)", line)
        if bullet:
            level = min(len(bullet.group(1)) // 2, 2)
            p = doc.add_paragraph()
            _apply_numbering(p, bullet_num_id, level)
            base.set_paragraph_spacing(p, before=0, after=4, line=1.25)
            base.add_inline_markdown(p, bullet.group(2), size=BODY_SIZE_PT, color=INK)
            index += 1
            continue

        p = base.add_para(doc, stripped, size=BODY_SIZE_PT, color=INK)
        if stripped.startswith("证据：") and len(stripped) <= 240:
            p.paragraph_format.keep_together = True
        index += 1

    if in_code:
        raise ValueError(f"Markdown 代码块在 EOF 前未闭合（起始行 L{code_start_line}）")
    _audit_table_geometry(doc)
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    print(build_docx())
